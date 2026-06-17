"""
Atualiza o Relatorio_ED2_241327.docx com os resultados finais dos 4 engines.
Uso: python scripts/update_report.py
"""
import json
import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


ROOT = Path(__file__).parent.parent


def load_summary(engine: str) -> dict:
    path = ROOT / "logs" / f"mec_summary_{engine}.json"
    with open(path) as f:
        return json.load(f)


def pct(val):
    if val is None:
        return "—"
    return f"{val*100:.1f}%"


def fmt(val, decimals=1):
    if val is None:
        return "—"
    return f"{val:.{decimals}f}"


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = text
    for run in cell.paragraphs[0].runs:
        run.bold = bold
        if color:
            run.font.color.rgb = RGBColor(*color)


def update_table_cell(table, row_idx, col_idx, text, bold=False, color=None):
    cell = table.cell(row_idx, col_idx)
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def find_and_replace_paragraph(doc, old_text, new_text):
    """Replace first occurrence of old_text in any paragraph."""
    for para in doc.paragraphs:
        if old_text in para.text:
            for run in para.runs:
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)
                    return True
    return False


def main():
    doc_path = ROOT / "Relatorio_ED2_241327.docx"
    if not doc_path.exists():
        print(f"ERROR: {doc_path} not found")
        sys.exit(1)

    print("Carregando resultados...")
    b = load_summary("BASELINE")
    d = load_summary("DRL")
    s = load_summary("SLM")
    l = load_summary("LLM")

    print(f"BASELINE: effective={b['effective_success_rate']:.1%}  anom={b.get('anomaly_compliance_rate','?')}")
    print(f"DRL:      effective={d['effective_success_rate']:.1%}  anom={d.get('anomaly_compliance_rate','?')}")
    print(f"SLM:      effective={s['effective_success_rate']:.1%}  anom={s.get('anomaly_compliance_rate','?')}")
    print(f"LLM:      effective={l['effective_success_rate']:.1%}  anom={l.get('anomaly_compliance_rate','?')}")

    doc = Document(str(doc_path))

    # ------------------------------------------------------------------ #
    # Table 0 — Engine overview (latência / energia)                     #
    # Row 2: DRL description update                                       #
    # ------------------------------------------------------------------ #
    t0 = doc.tables[0]
    # DRL row: "Q-learning tabular online" → "SB3 DQN (pré-treinado, 50k steps)"
    t0.cell(2, 1).text = "SB3 DQN (pré-treinado, 50k steps)"
    t0.cell(2, 2).text = "~0.6 ms (inferência MLP)"
    # SLM row: update model name
    t0.cell(3, 1).text = "Edge AI embarcada (Gemma 4 26B)"
    # LLM row: update model name
    t0.cell(4, 1).text = "Cloud AI centralizada (Gemini 3.1 Flash Lite)"
    t0.cell(4, 2).text = "~1-6 s (round-trip, inclui rate-limit)"

    # ------------------------------------------------------------------ #
    # Table 3 — KPI comparison (main results table)                       #
    # ------------------------------------------------------------------ #
    t3 = doc.tables[3]

    # Header row
    t3.cell(0, 2).text = "DRL (SB3)"
    t3.cell(0, 3).text = "SLM (Gemma 4)"
    t3.cell(0, 4).text = "LLM (Gemini 3.1)"

    # Row 1: Total tasks
    t3.cell(1, 2).text = str(d["total_tasks"])
    t3.cell(1, 3).text = str(s["total_tasks"])
    t3.cell(1, 4).text = str(l["total_tasks"])

    # Row 2: Success rate (throughput)
    t3.cell(2, 1).text = pct(b["success_rate"])
    t3.cell(2, 2).text = pct(d["success_rate"])
    t3.cell(2, 3).text = pct(s["success_rate"])
    t3.cell(2, 4).text = pct(l["success_rate"])

    # Row 3: Drop rate
    t3.cell(3, 1).text = pct(b["drop_rate"])
    t3.cell(3, 2).text = pct(d["drop_rate"])
    t3.cell(3, 3).text = pct(s["drop_rate"])
    t3.cell(3, 4).text = pct(l["drop_rate"])

    # Row 4: Latency
    t3.cell(4, 2).text = f"{fmt(d['avg_latency_ms'], 2)} ms"
    t3.cell(4, 3).text = "50 ms (NPU)"
    t3.cell(4, 4).text = f"{fmt(l['avg_latency_ms'])} ms"

    # Row 5: Total energy
    t3.cell(5, 1).text = f"{fmt(b['total_joules'], 3)} J"
    t3.cell(5, 2).text = f"{fmt(d['total_joules'], 3)} J"
    t3.cell(5, 3).text = f"{fmt(s['total_joules'], 1)} J"
    t3.cell(5, 4).text = f"{fmt(l['total_joules'], 1)} J"

    # Row 6: Energy/decision — fixed constants, no change needed

    # Row 7: Compliance semântico → Effective success rate
    t3.cell(7, 0).text = "Eff. taxa de sucesso"
    t3.cell(7, 1).text = pct(b["effective_success_rate"])
    t3.cell(7, 2).text = pct(d["effective_success_rate"])
    t3.cell(7, 3).text = pct(s["effective_success_rate"])
    t3.cell(7, 4).text = pct(l["effective_success_rate"])

    # Row 8: Anomaly compliance
    t3.cell(8, 1).text = pct(b.get("anomaly_compliance_rate"))
    t3.cell(8, 2).text = pct(d.get("anomaly_compliance_rate"))
    t3.cell(8, 3).text = pct(s.get("anomaly_compliance_rate"))
    t3.cell(8, 4).text = pct(l.get("anomaly_compliance_rate"))

    # Row 9: RAM utilization
    t3.cell(9, 1).text = f"{b.get('avg_ram_utilization_pct', 0):.1f}%"
    t3.cell(9, 2).text = f"{d.get('avg_ram_utilization_pct', 0):.1f}%"
    t3.cell(9, 3).text = f"{s.get('avg_ram_utilization_pct', 0):.1f}%"
    t3.cell(9, 4).text = f"{l.get('avg_ram_utilization_pct', 0):.1f}%"

    # ------------------------------------------------------------------ #
    # Table 4 — Battery states (BASELINE + other engines)                 #
    # ------------------------------------------------------------------ #
    t4 = doc.tables[4]
    # Add DRL, SLM, LLM final battery to caption or extend table
    sat1_drl  = d.get("sat_1_final_battery_pct", "?")
    sat2_drl  = d.get("sat_2_final_battery_pct", "?")
    sat15_drl = d.get("sat_15_final_battery_pct", "?")
    sat1_slm  = s.get("sat_1_final_battery_pct", "?")
    sat2_slm  = s.get("sat_2_final_battery_pct", "?")
    sat15_slm = s.get("sat_15_final_battery_pct", "?")
    sat1_llm  = l.get("sat_1_final_battery_pct", "?")
    sat2_llm  = l.get("sat_2_final_battery_pct", "?")
    sat15_llm = l.get("sat_15_final_battery_pct", "?")

    # Rebuild header to include all engines
    t4.cell(0, 3).text = "SOC Final BASELINE"
    # If table has only 4 cols, add info via paragraph replacement below

    # ------------------------------------------------------------------ #
    # Text paragraph updates                                              #
    # ------------------------------------------------------------------ #
    find_and_replace_paragraph(doc, "Gemma 3n via Gemini API", "Gemma 4 26B via Gemini API")
    find_and_replace_paragraph(doc, "Gemini 2.0 Flash Lite", "Gemini 3.1 Flash Lite")
    find_and_replace_paragraph(doc, "Q-learning tabular online", "SB3 DQN (pré-treinado, 50k steps)")
    find_and_replace_paragraph(doc, "Q-learning tabular real", "SB3 DQN (Stable Baselines3, 50k steps)")
    find_and_replace_paragraph(doc, "Q-learning tabular", "SB3 DQN")

    # Section 5.3 — update compliance numbers
    find_and_replace_paragraph(doc,
        "SLM (dados da arquitetura anterior) mostrou 37.5% de compliance em anomalias — s",
        f"SLM (Gemma 4 26B) obteve {pct(s.get('anomaly_compliance_rate'))} de compliance em anomalias — s")

    # Remove "trabalhos futuros" item about re-running SLM/LLM
    for para in doc.paragraphs:
        if "Re-execução completa de SLM e LLM" in para.text:
            para.text = ("✓ Concluído: SLM e LLM re-executados com arquitetura final "
                         "(eclipse orbital, ISL ideal, anomalias semânticas completas).")
            break

    # Section 7: Add note about current engine versions
    for para in doc.paragraphs:
        if "7. Exemplo de Output" in para.text:
            # Add new paragraph after section heading
            break

    # ------------------------------------------------------------------ #
    # Save                                                                #
    # ------------------------------------------------------------------ #
    out_path = ROOT / "Relatorio_ED2_241327.docx"
    doc.save(str(out_path))
    print(f"\nRelatório atualizado: {out_path}")

    # Print summary for verification
    print("\n=== SUMÁRIO DOS RESULTADOS ===")
    engines = [("BASELINE", b), ("DRL (SB3)", d), ("SLM (Gemma 4)", s), ("LLM (Gemini 3.1)", l)]
    print(f"{'Engine':<20} {'Success':>8} {'Effective':>10} {'AnomalyCmp':>12} {'Latency':>10}")
    print("-" * 65)
    for name, r in engines:
        print(f"{name:<20} {pct(r['success_rate']):>8} {pct(r['effective_success_rate']):>10} "
              f"{pct(r.get('anomaly_compliance_rate')):>12} {fmt(r['avg_latency_ms']):>8}ms")


if __name__ == "__main__":
    main()
