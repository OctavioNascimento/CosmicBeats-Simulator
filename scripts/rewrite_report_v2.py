"""
Reescrita completa do Relatorio_ED2_241327.docx: cenário/motivação mais claros,
métricas de compliance como KPI principal, e todos os números finais desta sessão
(DRL com reward shaping; SLM com o bug de parsing de 'thought parts' corrigido).

Uso: python scripts/rewrite_report_v2.py
"""
import json
from pathlib import Path
from docx import Document

ROOT = Path(__file__).parent.parent


def load_summary(engine):
    with open(ROOT / "logs" / f"mec_summary_{engine}.json") as f:
        return json.load(f)


def set_text(doc, index, new_text, label=""):
    p = doc.paragraphs[index]
    for run in p.runs:
        run.text = ""
    if p.runs:
        p.runs[0].text = new_text
    else:
        p.add_run(new_text)
    print(f"  P{index:<3} ✓ {label}")


def main():
    doc_path = ROOT / "Relatorio_ED2_241327.docx"
    doc = Document(str(doc_path))

    b = load_summary("BASELINE")
    d = load_summary("DRL")
    s = load_summary("SLM")
    l = load_summary("LLM")

    print("=== Reescrevendo Seção 1 (Introdução/Cenário) ===")
    set_text(doc, 5,
        "Cenário: uma constelação de três satélites em órbita baixa (LEO) atua "
        "como nós de computação de borda (MEC), cada um responsável por uma "
        "região geográfica (Brasil, Europa, EUA). Tarefas de processamento "
        "chegam continuamente e cada satélite decide, sozinho e em tempo real, "
        "se aceita, redireciona para outro satélite ou descarta a tarefa — sob "
        "restrições reais de bateria (ciclo solar/eclipse orbital) e RAM.",
        "cenário")
    set_text(doc, 6,
        "Para enriquecer o cenário, 10% das tarefas trazem uma restrição "
        "semântica do mundo real: uma lei de privacidade europeia (GDPR), uma "
        "exigência de soberania de dados do Brasil, ou uma falha crítica de "
        "hardware. Cumprir essas regras exige entender o significado da "
        "restrição, não apenas otimizar recursos — esse é o núcleo do "
        "experimento.",
        "motivação do cenário")
    set_text(doc, 8,
        "Hipótese central: orquestradores que leem e interpretam regras em "
        "linguagem natural (SLM embarcado e LLM na nuvem) tomam decisões mais "
        "corretas que uma heurística fixa (BASELINE) ou um agente de "
        "aprendizado por reforço puramente numérico (DRL), pagando em troca "
        "mais latência e energia. O valor científico do experimento é "
        "quantificar esse trade-off entre correção semântica e custo "
        "computacional — uma decisão de engenharia real para sistemas "
        "NTN-MEC, que operam sob orçamento de energia e banda extremamente "
        "restrito.",
        "valor científico")

    print("\n=== Reescrevendo Seção 2 (Adaptações) ===")
    set_text(doc, 10,
        "O plano original previa DRL via Stable Baselines3 (SB3) com Gym "
        "Wrapper — implementado com sucesso. Três componentes evoluíram além "
        "do plano original durante a implementação: o SLM, o modelo de "
        "bateria, e a premissa de conectividade entre satélites (ISL).",
        "2 intro")
    set_text(doc, 12,
        "O SB3 não pode ser conectado diretamente ao loop de eventos "
        "discretos do SimPy (que mantém estado compartilhado entre "
        "satélites), então o treinamento foi separado da execução: um "
        "ambiente Gymnasium sintético (NTNMECEnv) gera 50.000 episódios "
        "cobrindo a distribuição real de estados da simulação; o modelo "
        "treinado roda apenas inferência determinística (~0.5 ms) dentro da "
        "simulação.",
        "2.1 p1")
    set_text(doc, 13,
        "Observação: 13 valores normalizados — região da tarefa, um sinal "
        "binário \"há anomalia\" e o estado de bateria/RAM/luz-solar dos 3 "
        "satélites. O agente nunca vê o TIPO da anomalia, só que ela existe — "
        "isso preserva a comparação justa com SLM/LLM, que leem a regra em "
        "texto. Durante o treino, a recompensa usa o subtipo real da anomalia "
        "(oculto, nunca exposto ao agente) para ensinar a melhor estratégia "
        "possível dado essa cegueira parcial.",
        "2.1 p2 (reward shaping)")
    set_text(doc, 15,
        "Sem hardware NPU disponível para teste, o SLM é simulado via API do "
        "Gemma 4 26B (modelo pequeno, classe de parâmetros equivalente a um "
        "SLM embarcado). A latência usada nos KPIs é fixada em 50 ms — "
        "benchmark de NPU satelital da literatura — independente da latência "
        "real de rede da API, que é maior e serve apenas para diagnóstico.",
        "2.2 SLM")
    set_text(doc, 17,
        "A bateria segue um modelo físico de eclipse orbital: a iluminação "
        "solar varia como uma senoide de período 90 minutos (órbita LEO); "
        "abaixo de um limiar a bateria entra em eclipse e drena por consumo "
        "de housekeeping, fora do eclipse ela recarrega.",
        "2.3 p1")
    set_text(doc, 19,
        "Os três satélites têm capacidades (75, 50 e 100 Wh) e fases "
        "orbitais diferentes, simulando hardware heterogêneo real. O eclipse "
        "ocupa ~53% de cada órbita (46 dos 90 min), criando pressão "
        "energética real; abaixo de 20% de carga, o satélite recusa novas "
        "tarefas.",
        "2.3 p2")
    set_text(doc, 21,
        "O link entre satélites (ISL) é assumido sempre disponível, premissa "
        "comum em constelações modernas com ISL óptico — elimina uma "
        "variável de qualidade de sinal fora do escopo deste estudo. Drops "
        "refletem exclusivamente esgotamento de bateria/RAM ou recusa "
        "semântica, nunca falha de transmissão.",
        "2.4 ISL")

    print("\n=== Reescrevendo Seção 3 (Arquitetura) ===")
    set_text(doc, 33,
        "Como detalhado na Seção 2.1, o DRL não conhece o tipo da anomalia — "
        "apenas que ela existe — o que limita sua compliance semântica em "
        "comparação a SLM/LLM, que leem a regra em texto.",
        "3.3 invariante DRL")
    set_text(doc, 38,
        "Agente SB3 DQN pré-treinado em 50.000 episódios sintéticos via "
        "NTNMECEnv (Gymnasium), com recompensa moldada por subtipo de "
        "anomalia oculto durante o treino (nunca exposto na observação). "
        "Observação: vetor de 13 features — região da tarefa (one-hot 3D), "
        "has_anomaly, bateria/RAM/solar dos 3 satélites. Ações: Discrete(4) "
        "→ SAT-1 | SAT-2 | SAT-15 | DROP. Inferência determinística na "
        "simulação: ~0.5 ms (MLP forward pass).",
        "3.4.2 DRL")

    print("\n=== Reescrevendo Seção 4 (Desafios) ===")
    set_text(doc, 44,
        "RAM sem recuperação: tarefas consumiam RAM permanentemente, "
        "esgotando os 4096 MB após ~8 tarefas. Corrigido liberando RAM 60s "
        "após a conclusão de cada tarefa (TASK_DURATION_S).",
        "desafio 1")
    set_text(doc, 45,
        "DRL sem aprendizado real: a versão inicial usava pesos fixos "
        "(0.6×bateria + 0.4×RAM), sem qualquer treinamento. Substituído pelo "
        "agente SB3 DQN descrito na Seção 2.1.",
        "desafio 2")
    set_text(doc, 46,
        "Latência do SLM media o tempo de rede da API (segundos), não o "
        "hardware-alvo. Corrigido para usar a latência de benchmark de NPU "
        "(50 ms) nos KPIs, mantendo o tempo de rede real apenas como "
        "diagnóstico.",
        "desafio 3")
    set_text(doc, 47,
        "O prompt do LLM não incluía as regras semânticas nem um limitador "
        "de taxa, causando estouro do rate limit da API. Unificado com o "
        "mesmo prompt e regras do SLM, com limitador de 10 RPM.",
        "desafio 4")
    set_text(doc, 48,
        "O modelo de bateria estático foi substituído pelo modelo físico de "
        "eclipse orbital (Seção 2.3), após constatar que ele criava drops "
        "artificiais não relacionados a recursos reais.",
        "desafio 5")
    set_text(doc, 49,
        "API do SLM retorna o \"pensamento\" (thought) e a resposta final em "
        "partes separadas; o código inicial lia apenas a primeira parte "
        "(quase sempre o raciocínio, não a decisão), gerando ~40% de falhas "
        "de parsing que pareciam instabilidade de rede. Corrigido para usar "
        "apenas as partes finais (thought=false), eliminando praticamente "
        "todas as falhas espúrias (de 14 para 2 drops por causas de "
        "infraestrutura, de 69 tarefas).",
        "desafio 6 (bug real do thought-parts)")

    print("\n=== Reescrevendo Seção 5 (Resultados) ===")
    set_text(doc, 52,
        "Os experimentos foram executados com seed determinístico "
        "(random.seed(0)), garantindo que todos os engines recebam a mesma "
        "sequência das 69 tarefas. A janela de simulação é de 20 minutos "
        "(1200 s), com steps de 5 s.",
        "5 intro")
    set_text(doc, 55,
        "Nota: o SLM apresentou 2 quedas residuais por instabilidade de rede "
        "da API (de 69 chamadas), mesmo após a correção do bug de parsing "
        "(Seção 4) — não afetam a compliance semântica, que trata tarefas "
        "sem anomalia como corretas independentemente de descarte.",
        "5.1 footnote")
    set_text(doc, 62,
        "As duas métricas centrais deste estudo são semantic_compliance_rate "
        "(fração de decisões corretas, incluindo descartes corretos) e "
        "anomaly_compliance_rate (a mesma métrica, isolada às 8 tarefas com "
        "restrição semântica explícita). Essas, não a taxa de sucesso bruta, "
        "medem \"inteligência\" neste experimento: uma tarefa nunca está "
        "\"perdida\" — o satélite sempre tenta executar ou redirecionar; o "
        "que importa é se a decisão (incluindo descartar) foi a correta.",
        "5.3 intro")
    set_text(doc, 63,
        f"BASELINE acerta {b['semantic_compliance_rate']:.1%} das decisões "
        f"no geral, mas só {b['anomaly_compliance_rate']:.1%} (2/8) nas "
        "tarefas anômalas — sorte geográfica, não raciocínio: ele nunca "
        f"soube que havia uma regra a seguir. DRL (SB3 DQN) sobe para "
        f"{d['semantic_compliance_rate']:.1%} geral e "
        f"{d['anomaly_compliance_rate']:.1%} (3/8) em anomalias — superando "
        "o BASELINE nas duas métricas após um ajuste de recompensa que o "
        "ensina a reagir melhor ao sinal binário \"há anomalia\" (Seção "
        f"2.1). Esse {d['anomaly_compliance_rate']:.1%} é o teto teórico: "
        "das 8 anomalias reais, só as 3 falhas de hardware podem ser "
        "identificadas corretamente com um único bit de informação — as "
        "outras exigiriam saber qual regra específica se aplica, informação "
        "que o DRL estruturalmente não recebe.",
        "5.3 BASELINE/DRL")
    set_text(doc, 64,
        f"SLM (Gemma 4 26B) atinge {s['semantic_compliance_rate']:.1%} de "
        f"compliance geral e {s['anomaly_compliance_rate']:.1%} (6/8) em "
        f"anomalias, com effective_success_rate de "
        f"{s['effective_success_rate']:.1%} — a diferença para 100% reflete "
        "2 quedas residuais por instabilidade de rede da API (Seção 4), não "
        "erro de raciocínio: quando o modelo responde, a decisão está "
        f"correta. LLM (Gemini 3.1 Flash Lite) atinge "
        f"{l['semantic_compliance_rate']:.1%} em ambas as métricas, "
        "confirmando a hipótese central: ler a regra em linguagem natural "
        "supera tanto a heurística fixa quanto o aprendizado por reforço "
        "puramente numérico, na proporção exata da complexidade semântica "
        "da tarefa.",
        "5.3 SLM/LLM")

    print("\n=== Reescrevendo Seção 6 (Conclusão) ===")
    set_text(doc, 71,
        "DRL com SB3 DQN pré-treinado em 50.000 episódios (NTNMECEnv "
        "Gymnasium) e recompensa moldada por subtipo de anomalia oculto "
        "durante o treino: o agente supera o BASELINE em compliance geral "
        f"({d['semantic_compliance_rate']:.1%} vs "
        f"{b['semantic_compliance_rate']:.1%}) e em anomalias "
        f"({d['anomaly_compliance_rate']:.1%} vs "
        f"{b['anomaly_compliance_rate']:.1%}), sem nunca ver o tipo da "
        "anomalia — preservando o invariante científico que justifica a "
        "comparação com SLM/LLM.",
        "6.1 DRL")
    set_text(doc, 81,
        "▪ Aprendizado direcionado por recompensa vs. regras explícitas: o "
        "reward shaping elevou o DRL "
        f"({d['semantic_compliance_rate']:.1%} compliance, "
        f"{d['anomaly_compliance_rate']:.1%} em anomalias) acima do "
        "BASELINE nas duas métricas — mas o teto é estrutural: sem ver o "
        "tipo da anomalia, o DRL não pode chegar a 100% como o LLM. SLM/LLM "
        "aplicam regras zero-shot lidas em texto, sem essa limitação de "
        "informação.",
        "6.2 trade-off DRL")

    print("\n=== Reescrevendo Tabelas ===")
    t0 = doc.tables[0]
    t0.rows[2].cells[2].text = f"~{d['avg_latency_ms']:.1f} ms (inferência MLP)"
    print("  Tabela 0 ✓ latência DRL")

    t2 = doc.tables[2]
    t2.rows[3].cells[2].text = ("Todos os engines, incluindo DRL — única "
        "regra que o DRL aprendeu corretamente via reward shaping "
        "(responsável pelos 37.5% de anomaly_compliance)")
    print("  Tabela 2 ✓ quem acerta (hardware)")

    t3 = doc.tables[3]
    rows_data = [
        ("KPI", "BASELINE", "DRL (SB3)", "SLM (Gemma 4)", "LLM (Gemini 3.1)"),
        ("Total de tarefas", f"{b['total_tasks']}", f"{d['total_tasks']}", f"{s['total_tasks']}", f"{l['total_tasks']}"),
        ("Compliance semântica", f"{b['semantic_compliance_rate']:.1%}", f"{d['semantic_compliance_rate']:.1%}", f"{s['semantic_compliance_rate']:.1%}", f"{l['semantic_compliance_rate']:.1%}"),
        ("Compliance em anomalias", f"{b['anomaly_compliance_rate']:.1%}", f"{d['anomaly_compliance_rate']:.1%}", f"{s['anomaly_compliance_rate']:.1%}", f"{l['anomaly_compliance_rate']:.1%}"),
        ("Taxa de sucesso (throughput)", f"{b['success_rate']:.1%}", f"{d['success_rate']:.1%}", f"{s['success_rate']:.1%}", f"{l['success_rate']:.1%}"),
        ("Drop rate", f"{b['drop_rate']:.1%}", f"{d['drop_rate']:.1%}", f"{s['drop_rate']:.1%}", f"{l['drop_rate']:.1%}"),
        ("Latência média", "~0 ms", f"{d['avg_latency_ms']:.2f} ms", "50 ms (NPU)", f"{l['avg_latency_ms']:.1f} ms"),
        ("Energia total", f"{b['total_joules']:.3f} J", f"{d['total_joules']:.3f} J", f"{s['total_joules']:.1f} J", f"{l['total_joules']:.1f} J"),
        ("Energia/decisão", f"{b['joules_per_decision']:.3f} J", f"{d['joules_per_decision']:.3f} J", f"{s['joules_per_decision']:.3f} J", f"{l['joules_per_decision']:.3f} J"),
        ("RAM utilização final", f"{b['avg_ram_utilization_pct']:.1f}%", f"{d['avg_ram_utilization_pct']:.1f}%", f"{s['avg_ram_utilization_pct']:.1f}%", f"{l['avg_ram_utilization_pct']:.1f}%"),
    ]
    for ri, row_vals in enumerate(rows_data):
        for ci, val in enumerate(row_vals):
            t3.rows[ri].cells[ci].text = val
    print("  Tabela 3 ✓ reescrita completa (KPIs reordenados, compliance primeiro)")

    doc.save(str(doc_path))
    print(f"\n✓ Relatório salvo: {doc_path}")


if __name__ == "__main__":
    main()
